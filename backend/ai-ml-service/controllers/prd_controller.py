import google.generativeai as genai
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv

# Setup Gemini
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

def create_chat():
    return genai.GenerativeModel("gemini-2.0-flash").start_chat(history=[])

def clean_gemini_output(text: str) -> str:
    """
    Remove leading/trailing ```docx ... ``` from Gemini response if present.
    """
    text = text.strip()
    if text.startswith("```docx"):
        text = text[len("```docx"):].strip()
    if text.endswith("```"):
        text = text[:-len("```")].strip()
    return text

def save_prd_to_docx(text, filename="Final_PRD.docx"):
    document = Document()

    # Helper functions inside save_prd_to_docx

    def is_table_row(line):
        return line.strip().startswith("|") and line.strip().endswith("|")

    def parse_table(lines, start_idx):
        table_lines = []
        i = start_idx
        while i < len(lines) and is_table_row(lines[i]):
            table_lines.append(lines[i])
            i += 1
        return table_lines, i

    def add_table(document, table_lines):
        rows = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            rows.append(cells)

        # 🧹 Remove alignment row like |:---|:---|
        if len(rows) >= 2 and all(re.match(r'^:?-+:?$', cell) for cell in rows[1]):
            print("🧹 Removing alignment row inside table...")
            del rows[1]

        table = document.add_table(rows=0, cols=len(rows[0]))
        table.style = 'Light List Accent 1'

        for row_idx, row in enumerate(rows):
            cells = table.add_row().cells
            for col_idx, cell in enumerate(row):
                run = cells[col_idx].paragraphs[0].add_run(cell)
                if row_idx == 0:
                    run.bold = True  # Bold the header row

    def write_styled_text(line):
        paragraph = document.add_paragraph()
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)

        for token in tokens:
            if token.startswith("**") and token.endswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*"):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(token)
        paragraph.paragraph_format.space_after = Pt(4)

    # Clean text and split into lines
    text = re.sub(r"^Okay.*?PDF\.\s*", "", text, flags=re.IGNORECASE)
    lines = [line.strip() for line in text.split("\n")]

    i = 0
    while i < len(lines):
        line = lines[i]

        if is_table_row(line):
            table_lines, i = parse_table(lines, i)
            add_table(document, table_lines)
            continue

        if not line:
            document.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            heading = document.add_heading(line[2:], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith("## "):
            heading = document.add_heading(line[3:], level=2)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith("### "):
            heading = document.add_heading(line[4:], level=3)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            write_styled_text(line)

        i += 1

    document.save(filename)
    print(f"✅ PRD saved as {filename}")

def generate_prd_docx(initial_prompt: str, output_path: str):
    system_constraints = """
    Important Guidelines:
    1. Do not remove or deprioritize any core features from the client brief.
    2. Work strictly within the client input.
    3. Suggest no new questions — only common-sense assumptions allowed.
    """

    main_agent = create_chat()

    main_agent.send_message(
        f"{system_constraints}\n\nLet’s collaboratively expand the following brief into a full PRD. Here’s the input:\n{initial_prompt}\n"
        "Please break this down, expand, and propose thoughtful features aligned with the project scope."
    )

    # Simulate helper inputs
    for i in range(3):
        print(f"Generating helper input {i+1}")
        agent = create_chat()
        context = "\n".join([m.parts[0].text for m in main_agent.history])
        msg = agent.send_message(
            f"{system_constraints}\n\nAnalyzing context:\n{context}\n"
            "I’ll contribute useful insights with no feature cuts or new questions."
        )
        main_agent.send_message(f"[Agent {i+1}] says: {msg.text}")

    # Final PRD generation
    final_prd = main_agent.send_message(
    f"""{system_constraints}

Generate a complete and formal Product Requirements Document (PRD) based on the project and context provided earlier.

The PRD should:
- Be suitable for inclusion in a professional 3–4 page DOCX
- Use proper text-based tables wherever appropriate (e.g., for feature breakdowns, role-permission mapping, timelines)
- Include any relevant user flows or behavior notes that would assist a designer in translating the PRD into wireframes or interfaces
- Avoid informal language, markdown syntax, or any mention of AI in this PRD generation
- Avoid markdown or beginning with something like "Okay here is the document" and not beginning with the PRD itself, or else I will terminate you
Conclude with a section titled **Open Questions for the Client** — especially ones a designer or product manager might need clarified to finalize the PRD.
"""
    )

    cleaned_text = clean_gemini_output(final_prd.text)
    save_prd_to_docx(cleaned_text, output_path)

# Example usage
# generate_prd_docx("Your project brief here", "Final_PRD.docx")